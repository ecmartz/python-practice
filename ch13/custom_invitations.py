#!/usr/bin/env python3

import os
import docx
from docx import text

guests_file = "./guests.txt"
guests_doc = "./guests.docx"

if __name__ == "__main__":
    document = docx.Document(guests_doc)
    num = 0
    with open(guests_file, 'r') as F:
        for line in F.readlines():
            print(line)
            document.add_paragraph("It would be a pleasure to have the company of")
            document.paragraphs[-1].style = 'Scribble'
            document.add_paragraph(line)
            document.paragraphs[-1].style = 'Big1'
            document.add_paragraph("at 11010 Memory Lane on the Evening of")
            document.paragraphs[-1].style = 'Scribble'
            document.add_paragraph("April 1st")
            document.paragraphs[-1].style = 'Big2'
            document.add_paragraph("at 7 o'clock")
            document.paragraphs[-1].style = 'Scribble'
            document.add_page_break()
    document.save(guests_doc)
